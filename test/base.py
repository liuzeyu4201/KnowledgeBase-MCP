from __future__ import annotations

import asyncio
import base64
import binascii
from io import BytesIO
import os
import time
import unittest
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from docx import Document as DocxDocument
import httpx
from minio import Minio
from minio.error import S3Error

from test.mcp_test_client import MCPToolClient


class MCPIntegrationTestCase(unittest.TestCase):
    """提供统一 MCP 测试入口、样例文件读取和测试辅助能力。"""

    server_url = os.getenv("KNOWLEDGEBASE_TEST_SERVER_URL", "http://127.0.0.1:8000/mcp")
    api_base_url = os.getenv("KNOWLEDGEBASE_TEST_API_BASE_URL", "http://127.0.0.1:8000")
    small_pdf_path = Path(
        os.getenv("KNOWLEDGEBASE_TEST_SMALL_PDF", "/app/data/Functional_Analysis.pdf")
    )
    large_pdf_path = Path(
        os.getenv("KNOWLEDGEBASE_TEST_LARGE_PDF", "/app/data/Functional Analysis Notes.pdf")
    )
    minio_endpoint = os.getenv(
        "KNOWLEDGEBASE_TEST_MINIO_ENDPOINT",
        os.getenv("KNOWLEDGEBASE_MINIO_ENDPOINT", "127.0.0.1:9000"),
    )
    minio_access_key = os.getenv(
        "KNOWLEDGEBASE_TEST_MINIO_ACCESS_KEY",
        os.getenv("KNOWLEDGEBASE_MINIO_ACCESS_KEY", "minioadmin"),
    )
    minio_secret_key = os.getenv(
        "KNOWLEDGEBASE_TEST_MINIO_SECRET_KEY",
        os.getenv("KNOWLEDGEBASE_MINIO_SECRET_KEY", "minioadmin"),
    )
    minio_secure = os.getenv(
        "KNOWLEDGEBASE_TEST_MINIO_SECURE",
        os.getenv("KNOWLEDGEBASE_MINIO_SECURE", "false"),
    ).lower() in {"1", "true", "yes", "on"}

    def run_async(self, coroutine):
        """在同步 unittest 用例中执行异步逻辑。"""

        return asyncio.run(coroutine)

    def unique_suffix(self) -> str:
        """生成唯一后缀，避免测试数据在共享环境中发生冲突。"""

        return f"{int(time.time() * 1000)}_{id(self)}"

    def read_pdf_base64(self, *, large: bool = False) -> str:
        """读取测试 PDF 并编码为 base64。"""

        pdf_path = self.large_pdf_path if large else self.small_pdf_path
        return base64.b64encode(pdf_path.read_bytes()).decode("utf-8")

    def read_pdf_bytes(self, *, large: bool = False) -> bytes:
        """读取测试 PDF 原始字节。"""

        pdf_path = self.large_pdf_path if large else self.small_pdf_path
        return pdf_path.read_bytes()

    def read_markdown_base64(self, *, title: str = "Markdown Title") -> str:
        """构造 Markdown 文档并编码为 base64。"""

        markdown = (
            f"# {title}\n\n"
            "这是一个 Markdown 导入测试文档。\n\n"
            "- 第一条\n"
            "- 第二条\n\n"
            "## 子标题\n\n"
            "线性代数与泛函分析会在这里相遇。"
        )
        return base64.b64encode(markdown.encode("utf-8")).decode("utf-8")

    def read_markdown_bytes(self, *, title: str = "Markdown Title") -> bytes:
        """构造 Markdown 文档原始字节。"""

        markdown = (
            f"# {title}\n\n"
            "这是一个 Markdown 导入测试文档。\n\n"
            "- 第一条\n"
            "- 第二条\n\n"
            "## 子标题\n\n"
            "线性代数与泛函分析会在这里相遇。"
        )
        return markdown.encode("utf-8")

    def read_docx_base64(self, *, title: str = "Docx Title") -> str:
        """构造 DOCX 文档并编码为 base64。"""

        document = DocxDocument()
        document.add_heading(title, level=1)
        document.add_paragraph("这是一个 Word 导入测试文档。")
        document.add_paragraph("它包含多段正文，用于验证解析、切片和向量写入链路。")
        document.add_paragraph("Hilbert space is a complete inner product space.")
        buffer = BytesIO()
        document.save(buffer)
        return base64.b64encode(buffer.getvalue()).decode("utf-8")

    def read_docx_bytes(self, *, title: str = "Docx Title") -> bytes:
        """构造 DOCX 文档原始字节。"""

        document = DocxDocument()
        document.add_heading(title, level=1)
        document.add_paragraph("这是一个 Word 导入测试文档。")
        document.add_paragraph("它包含多段正文，用于验证解析、切片和向量写入链路。")
        document.add_paragraph("Hilbert space is a complete inner product space.")
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    async def tool(self, tool_name: str, **arguments: Any) -> dict[str, Any]:
        """统一调用 MCP Tool。"""

        async with MCPToolClient(self.server_url) as client:
            return await client.call_tool(tool_name, arguments)

    def http_json(
        self,
        method: str,
        path: str,
        *,
        params: dict[str, Any] | None = None,
        files: dict[str, Any] | None = None,
        data: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """调用普通 HTTP 接口并返回 JSON。"""

        response = httpx.request(
            method,
            f"{self.api_base_url}{path}",
            params=params,
            files=files,
            data=data,
            timeout=120.0,
        )
        return response.json()

    def http_text(
        self,
        method: str,
        path: str,
        *,
        params: dict[str, Any] | None = None,
    ) -> tuple[int, str]:
        """调用普通 HTTP 接口并返回状态码和文本内容。"""

        response = httpx.request(
            method,
            f"{self.api_base_url}{path}",
            params=params,
            timeout=120.0,
        )
        return response.status_code, response.text

    def assert_success(self, payload: dict[str, Any], *, code: str = "OK") -> None:
        """断言接口成功返回。"""

        self.assertTrue(payload.get("success"), payload)
        self.assertEqual(payload.get("code"), code, payload)

    def assert_error(
        self,
        payload: dict[str, Any],
        *,
        code: str | None = None,
        error_type: str | None = None,
    ) -> None:
        """断言接口失败返回，并可校验错误码与错误类型。"""

        self.assertFalse(payload.get("success"), payload)
        if code is not None:
            self.assertEqual(payload.get("code"), code, payload)
        if error_type is not None:
            self.assertEqual(payload.get("error", {}).get("type"), error_type, payload)

    async def create_category(self, *, prefix: str = "test_category") -> dict[str, Any]:
        """创建分类并登记清理信息。"""

        suffix = self.unique_suffix()
        payload = await self.tool(
            "kb_category_create",
            category_code=f"{prefix}_{suffix}",
            name=f"{prefix}_{suffix}",
            description="测试分类",
        )
        self.assert_success(payload)
        return payload["data"]["category"]

    async def import_document(
        self,
        *,
        category_id: int,
        title_prefix: str = "test_document",
        large: bool = False,
        file_name: str | None = None,
        mime_type: str = "application/pdf",
        file_content_base64: str | None = None,
    ) -> dict[str, Any]:
        """通过 staged_file 标准路径导入测试文档。"""

        suffix = self.unique_suffix()
        resolved_file_name = file_name or ("Functional Analysis Notes.pdf" if large else "Functional_Analysis.pdf")
        resolved_file_content_base64 = file_content_base64 or self.read_pdf_base64(large=large)
        staged_file = self.upload_staged_file(
            file_name=resolved_file_name,
            file_bytes=base64.b64decode(resolved_file_content_base64),
            mime_type=mime_type,
        )
        payload = await self.tool(
            "kb_document_import_from_staged",
            category_id=category_id,
            title=f"{title_prefix}_{suffix}",
            staged_file_id=staged_file["id"],
        )
        document, _ = await self.resolve_document_write_payload(payload)
        return document

    async def update_document(
        self,
        *,
        document_id: int,
        category_id: int | None = None,
        title: str | None = None,
        file_name: str | None = None,
        mime_type: str | None = None,
        file_content_base64: str | None = None,
    ) -> dict[str, Any]:
        """通过 staged_file 标准路径更新文档。"""

        payload: dict[str, Any] = {
            "id": document_id,
            "category_id": category_id,
            "title": title,
        }
        if file_content_base64 is not None:
            if file_name is None or mime_type is None:
                raise AssertionError("file_name 和 mime_type 必须与 file_content_base64 一起提供")
            staged_file = self.upload_staged_file(
                file_name=file_name,
                file_bytes=base64.b64decode(file_content_base64),
                mime_type=mime_type,
            )
            payload["staged_file_id"] = staged_file["id"]

        result = await self.tool("kb_document_update_from_staged", **payload)
        if result.get("success") and "task" in result.get("data", {}):
            document, task = await self.resolve_document_write_payload(
                result,
                document_id_hint=document_id,
            )
            return {
                **result,
                "data": {
                    "document": document,
                    "task": task,
                },
            }
        return result

    async def submit_batch_import_task(
        self,
        *,
        items: list[dict[str, Any]],
        priority: int = 50,
        max_attempts: int = 3,
        idempotency_key: str | None = None,
        request_id: str | None = None,
        operator: str | None = None,
        trace_id: str | None = None,
    ) -> dict[str, Any]:
        """测试辅助：把文件项先上传为 staged_file，再调用标准批量任务 Tool。"""

        normalized_items: list[dict[str, Any]] = []
        for item in items:
            if item.get("staged_file_id") is not None:
                normalized_items.append(
                    {
                        "category_id": item["category_id"],
                        "title": item["title"],
                        "staged_file_id": item["staged_file_id"],
                        "priority": item.get("priority"),
                    }
                )
                continue

            file_name = item.get("file_name")
            if not isinstance(file_name, str) or not file_name.strip() or len(file_name.strip()) > 256:
                return self._build_validation_error_payload("INVALID_ARGUMENT", "validation_error", "file_name 不合法")

            file_content_base64 = item.get("file_content_base64")
            if not isinstance(file_content_base64, str) or not file_content_base64:
                return self._build_validation_error_payload("INVALID_ARGUMENT", "validation_error", "file_content_base64 不合法")

            try:
                file_bytes = base64.b64decode(file_content_base64, validate=True)
            except (ValueError, binascii.Error):
                return self._build_validation_error_payload("INVALID_ARGUMENT", "validation_error", "file_content_base64 不合法")

            mime_type = item.get("mime_type")
            if not isinstance(mime_type, str) or not mime_type.strip():
                return self._build_validation_error_payload("INVALID_ARGUMENT", "validation_error", "mime_type 不合法")
            upload_payload = self.http_json(
                "POST",
                "/api/staged-files",
                files={
                    "file": (file_name, file_bytes, mime_type),
                },
            )
            if not upload_payload.get("success"):
                return upload_payload

            staged_file = upload_payload["data"]["staged_file"]
            normalized_items.append(
                {
                    "category_id": item["category_id"],
                    "title": item["title"],
                    "staged_file_id": staged_file["id"],
                    "priority": item.get("priority"),
                }
            )

        return await self.tool(
            "kb_document_import_batch_submit_from_staged",
            items=normalized_items,
            priority=priority,
            max_attempts=max_attempts,
            idempotency_key=idempotency_key,
            request_id=request_id,
            operator=operator,
            trace_id=trace_id,
        )

    def upload_staged_file(
        self,
        *,
        file_name: str,
        file_bytes: bytes,
        mime_type: str,
    ) -> dict[str, Any]:
        """通过普通 HTTP 上传接口创建暂存文件。"""

        payload = self.http_json(
            "POST",
            "/api/staged-files",
            files={
                "file": (file_name, file_bytes, mime_type),
            },
        )
        self.assert_success(payload)
        return payload["data"]["staged_file"]

    def _build_validation_error_payload(
        self,
        code: str,
        error_type: str,
        message: str,
    ) -> dict[str, Any]:
        """为测试辅助函数构造统一错误响应。"""

        return {
            "success": False,
            "code": code,
            "message": message,
            "error": {
                "type": error_type,
                "details": {},
            },
        }

    async def delete_document_best_effort(self, document: dict[str, Any]) -> None:
        """尽最大努力删除测试文档，避免污染共享测试环境。"""

        await self.tool("kb_document_delete", id=document["id"])

    async def delete_category_best_effort(self, category: dict[str, Any]) -> None:
        """尽最大努力删除测试分类。"""

        await self.tool("kb_category_delete", id=category["id"])

    async def delete_documents_by_category_best_effort(self, *, category_id: int) -> None:
        """尽最大努力删除某个分类下的全部文档。"""

        payload = await self.tool(
            "kb_document_list",
            category_id=category_id,
            page=1,
            page_size=100,
        )
        if not payload.get("success"):
            return
        for item in payload["data"]["items"]:
            await self.tool("kb_document_delete", id=item["id"])

    async def resolve_document_write_payload(
        self,
        payload: dict[str, Any],
        *,
        document_id_hint: int | None = None,
        timeout_seconds: float = 120.0,
    ) -> tuple[dict[str, Any], dict[str, Any] | None]:
        """把文档写入返回统一解析成最终文档。"""

        self.assert_success(payload)
        data = payload.get("data", {})
        document = data.get("document")
        if isinstance(document, dict):
            return document, None

        task = data.get("task")
        if not isinstance(task, dict):
            self.fail(f"unexpected document write payload: {payload}")

        return await self.wait_for_document_task_document(
            task=task,
            document_id_hint=document_id_hint,
            timeout_seconds=timeout_seconds,
        )

    async def wait_for_document_task_document(
        self,
        *,
        task: dict[str, Any],
        document_id_hint: int | None = None,
        timeout_seconds: float = 120.0,
    ) -> tuple[dict[str, Any], dict[str, Any]]:
        """轮询文档异步任务，直到拿到最终文档。"""

        deadline = time.monotonic() + timeout_seconds
        final_task = task
        while time.monotonic() < deadline:
            get_payload = await self.tool(
                "kb_document_task_get",
                id=task["id"],
                include_items=True,
            )
            self.assert_success(get_payload)
            final_task = get_payload["data"]["task"]
            if final_task["status"] in {"success", "partial_success", "failed", "canceled"}:
                break
            await asyncio.sleep(1)

        self.assertIn(final_task["status"], {"success", "partial_success"}, final_task)
        items = final_task.get("items") or []
        document_id = document_id_hint
        if items:
            document_id = items[0].get("document_id") or document_id
        self.assertIsNotNone(document_id, final_task)

        get_payload = await self.tool("kb_document_get", id=document_id)
        self.assert_success(get_payload)
        return get_payload["data"]["document"], final_task

    def assert_storage_uri_exists(self, storage_uri: str) -> None:
        """断言对象或本地文件存在。"""

        if storage_uri.startswith("s3://"):
            bucket, object_key = self.parse_s3_uri(storage_uri)
            client = self.build_minio_client()
            stat = client.stat_object(bucket, object_key)
            self.assertIsNotNone(stat)
            return
        self.assertTrue(Path(storage_uri).exists(), storage_uri)

    def wait_for_storage_uri_deleted(self, storage_uri: str, *, timeout_seconds: float = 10.0) -> None:
        """轮询等待对象或本地文件被删除。"""

        deadline = time.time() + timeout_seconds
        while time.time() < deadline:
            if not self.storage_uri_exists(storage_uri):
                return
            time.sleep(0.2)
        self.fail(f"storage uri still exists after timeout: {storage_uri}")

    def storage_uri_exists(self, storage_uri: str) -> bool:
        """判断对象或本地文件是否存在。"""

        if storage_uri.startswith("s3://"):
            bucket, object_key = self.parse_s3_uri(storage_uri)
            client = self.build_minio_client()
            try:
                client.stat_object(bucket, object_key)
                return True
            except S3Error as exc:
                if exc.code in {"NoSuchKey", "NoSuchObject", "NoSuchBucket"}:
                    return False
                raise
        return Path(storage_uri).exists()

    def build_minio_client(self) -> Minio:
        """构造测试用 MinIO 客户端。"""

        return Minio(
            self.minio_endpoint,
            access_key=self.minio_access_key,
            secret_key=self.minio_secret_key,
            secure=self.minio_secure,
        )

    def parse_s3_uri(self, storage_uri: str) -> tuple[str, str]:
        """解析 s3://bucket/key 形式的存储 URI。"""

        parsed = urlparse(storage_uri)
        if parsed.scheme != "s3":
            raise AssertionError(f"unsupported storage uri: {storage_uri}")
        return parsed.netloc, parsed.path.lstrip("/")
